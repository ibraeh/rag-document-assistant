"""
Document processing service for extracting and chunking text
RAG Document Assistant - Complete Document Processor
"""
import os
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import logging

import pypdf
from docx import Document as DocxDocument
# from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import settings

logger = logging.getLogger(__name__)


# ============================================================================
# CHUNKING STRATEGIES
# ============================================================================

class ChunkingStrategy:
    @staticmethod
    def semantic_chunking(text: str, chunk_size: int, overlap: int) -> List[str]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len,
            separators=["\\n\\n", "\\n", ". ", "! ", "? ", "; ", ", ", " ", ""]
        )
        return splitter.split_text(text)

    @staticmethod
    def fixed_chunking(text: str, chunk_size: int, overlap: int) -> List[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if chunk.strip():
                chunks.append(chunk)
            start += chunk_size - overlap
        return chunks

    @staticmethod
    def paragraph_chunking(text: str, max_chunk_size: int) -> List[str]:
        paragraphs = text.split('\\n\\n')
        chunks, current_chunk, current_size = [], [], 0
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if current_size + len(para) > max_chunk_size and current_chunk:
                chunks.append('\\n\\n'.join(current_chunk))
                current_chunk = [para]
                current_size = len(para)
            else:
                current_chunk.append(para)
                current_size += len(para)
        if current_chunk:
            chunks.append('\\n\\n'.join(current_chunk))
        return chunks


# ============================================================================
# TEXT CLEANING
# ============================================================================

class TextCleaner:
    @staticmethod
    def clean_text(text: str) -> str:
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\\n{3,}', '\\n\\n', text)
        text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\'\n]", "", text)
        return text.strip()

    @staticmethod
    def remove_headers_footers(text: str, patterns: Optional[List[str]] = None) -> str:
        patterns = patterns or [
            r'Page \\d+ of \\d+',
            r'\\d+\\s*$',
            r'^Copyright ©.*$',
            r'^All rights reserved.*$'
        ]
        for pattern in patterns:
            text = re.sub(pattern, '', text, flags=re.MULTILINE | re.IGNORECASE)
        return text

    @staticmethod
    def normalize_whitespace(text: str) -> str:
        text = text.replace('\\t', ' ').replace('\\r', '')
        text = re.sub(r' +', ' ', text)
        return '\\n'.join(line.strip() for line in text.split('\\n'))


# ============================================================================
# DOCUMENT PROCESSOR
# ============================================================================

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\\n\\n", "\\n", ". ", "! ", "? ", " ", ""]
        )
        self.text_cleaner = TextCleaner()
        self.chunking_strategy = ChunkingStrategy()

    def generate_document_id(self, filename: str, content: bytes) -> str:
        content_hash = hashlib.md5(content).hexdigest()[:8]
        clean_name = re.sub(r'[^\w\\-]', '_', Path(filename).stem)[:30]
        timestamp = datetime.now().strftime("%Y%m%d")
        return f"{clean_name}_{timestamp}_{content_hash}"

    def extract_text(self, file_path: str, file_type: str) -> Tuple[str, int, Dict]:
        if file_type == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif file_type == ".docx":
            return self.extract_text_from_docx(file_path)
        elif file_type == ".txt":
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, int, Dict]:
        try:
            text_parts, metadata = [], {'page_sizes': []}
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                num_pages = len(reader.pages)
                if reader.metadata:
                    metadata.update({
                        'title': reader.metadata.get('/Title'),
                        'author': reader.metadata.get('/Author'),
                        'creation_date': reader.metadata.get('/CreationDate')
                    })
                for i, page in enumerate(reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page.mediabox:
                            metadata['page_sizes'].append({
                                'width': float(page.mediabox.width),
                                'height': float(page.mediabox.height)
                            })
                        if page_text and page_text.strip():
                            cleaned = self.text_cleaner.normalize_whitespace(page_text)
                            text_parts.append(f"\\n[Page {i}]\\n{cleaned}\\n")
                        else:
                            text_parts.append(f"\\n[Page {i}]\\n[No extractable text]\\n")
                    except Exception as e:
                        logger.warning(f"Page {i} error: {e}")
                        text_parts.append(f"\\n[Page {i}]\\n[Extraction error]\\n")
            full_text = ''.join(text_parts)
            if len(full_text.strip()) < 50:
                raise ValueError("PDF text too short or unreadable")
            return full_text, num_pages, metadata
        except Exception as e:
            raise ValueError(f"PDF extraction failed: {e}")

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, int, Dict]:
        try:
            doc = DocxDocument(file_path)
            metadata = {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text_parts.append(row_text)
            full_text = self.text_cleaner.clean_text(self.text_cleaner.normalize_whitespace('\\n\\n'.join(text_parts)))
            word_count = len(full_text.split())
            estimated_pages = max(1, word_count // 500)
            if len(full_text.strip()) < 50:
                raise ValueError("DOCX text too short")
            return full_text, estimated_pages, metadata
        except Exception as e:
            raise ValueError(f"DOCX extraction failed: {e}")

    def extract_text_from_txt(self, file_path: str) -> Tuple[str, int, Dict]:
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text, encoding_used = None, None
            for enc in encodings:
                try:
                    with open(file_path, 'r', encoding=enc) as f:
                        text = f.read()
                    encoding_used = enc
                    break
                except UnicodeDecodeError:
                    continue
            if text is None:
                raise ValueError("TXT decoding failed")
            text = self.text_cleaner.normalize_whitespace(text)
            word_count = len(text.split())
            estimated_pages = max(1, word_count // 500)
            metadata = {
                'encoding': encoding_used,
                'size_bytes': os.path.getsize(file_path),
                'lines': len(text.split('\\n')),
                'words': word_count
            }
            if len(text.strip()) < 10:
                raise ValueError("TXT file too short")
            return text, estimated_pages, metadata
        except Exception as e:
            raise ValueError(f"TXT extraction failed: {e}")

    def chunk_text(self, text: str, metadata: Dict, strategy: str = "semantic") -> List[Dict]:
        if not text or len(text.strip()) < 10:
            raise ValueError("Text too short to chunk")
        if strategy == "semantic":
            chunks = self.chunking_strategy.semantic_chunking(text, settings.chunk_size, settings.chunk_overlap)
        elif strategy == "fixed":
            chunks = self.chunking_strategy.fixed_chunking(text, settings.chunk_size, settings.chunk_overlap)
        elif strategy == "paragraph":
            chunks = self.chunking_strategy.paragraph_chunking(text, settings.chunk_size)
        else:
            raise ValueError(f"Unknown chunking strategy: {strategy}")

        result = []
        for i, chunk_text in enumerate(chunks):
            if not chunk_text.strip():
                continue
            page_num = self._extract_page_number(chunk_text)
            result.append({
                "text": chunk_text.strip(),
                "metadata": {
                    **metadata,
                    "chunk_id": i,
                    "chunk_index": i,
                    "page_number": page_num,
                    "chunk_size": len(chunk_text),
                    "chunk_strategy": strategy
                }
            })

        if not result:
            raise ValueError("No valid chunks created from text")

        logger.info(f"Created {len(result)} chunks using {strategy} strategy")
        return result

    def _extract_page_number(self, text: str) -> int:
        match = re.search(r'\[Page (\d+)\]', text)
        return int(match.group(1)) if match else 1

    def process_document(self, file_path: str, filename: str, file_type: str, chunking_strategy: str = "semantic", clean_text: bool = True) -> Dict:
        try:
            start_time = datetime.now()
            text, num_pages, extraction_metadata = self.extract_text(file_path, file_type)

            if clean_text:
                text = self.text_cleaner.remove_headers_footers(text)
                text = self.text_cleaner.clean_text(text)

            if len(text.strip()) < 50:
                raise ValueError("Extracted text is too short or empty")

            with open(file_path, "rb") as f:
                content = f.read()
            document_id = self.generate_document_id(filename, content)

            metadata = {
                "document_id": document_id,
                "filename": filename,
                "file_type": file_type,
                "pages": num_pages,
                "source": file_path,
                "processed_at": datetime.now().isoformat(),
                "file_size": len(content),
                "text_length": len(text),
                "word_count": len(text.split())
            }

            chunks = self.chunk_text(text, metadata, strategy=chunking_strategy)
            elapsed = (datetime.now() - start_time).total_seconds()

            return {
                "document_id": document_id,
                "text": text,
                "chunks": chunks,
                "pages": num_pages,
                "metadata": metadata,
                "extraction_metadata": extraction_metadata,
                "processing_time": elapsed,
                "chunk_count": len(chunks)
            }
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise

    def validate_file(self, filename: str, file_size: int) -> None:
        ext = Path(filename).suffix.lower()
        if ext not in settings.allowed_extensions:
            raise ValueError(f"File type '{ext}' not allowed. Supported formats: {', '.join(settings.allowed_extensions)}")
        if file_size > settings.max_file_size:
            raise ValueError(f"File too large ({file_size / 1024 / 1024:.1f}MB). Max size: {settings.max_file_size / 1024 / 1024:.1f}MB")
        if file_size < 100:
            raise ValueError("File is too small or empty")
        logger.info(f"File validation passed: {filename} ({file_size} bytes)")

    def get_supported_formats(self) -> List[str]:
        return list(settings.allowed_extensions)

    def estimate_processing_time(self, file_size: int, file_type: str) -> float:
        rates = {
            ".pdf": 500_000,
            ".docx": 1_000_000,
            ".txt": 2_000_000
        }
        rate = rates.get(file_type, 500_000)
        return file_size / rate


# ============================================================================
# SINGLETON INSTANCE
# ============================================================================

document_processor = DocumentProcessor()


# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def process_file(file_path: str, filename: str) -> Dict:
    file_type = Path(filename).suffix.lower()
    return document_processor.process_document(file_path, filename, file_type)

def validate_upload(filename: str, file_size: int):
    document_processor.validate_file(filename, file_size)


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    print("\n=== Document Processor Test ===\n")

    try:
        document_processor.validate_file("test.pdf", 1024000)
        print("✓ Validation passed for test.pdf")
    except ValueError as e:
        print(f"✗ Validation failed: {e}")

    print(f"\nSupported formats: {document_processor.get_supported_formats()}")
    est_time = document_processor.estimate_processing_time(2048000, ".pdf")
    print(f"\nEstimated processing time for 2MB PDF: {est_time:.2f}s")
    print("\n=== Test Complete ===\n")



